"""Job-description file text extraction.

Supported per instructions.txt: PDF, DOC, DOCX, TXT, HTML. Each parser returns
plain text with paragraph structure preserved (blank-line separated), which is
what the downstream stripping/normalization LLM calls consume.

Encoding note: client-supplied files need two separate defences, and the real
sample JDs in `Legacy jaStudio/2. Job Profile/a.Before/` exercise both:
  1. Windows-authored .txt is often cp1252 rather than UTF-8 — hence encoding
     sniffing rather than assuming UTF-8.
  2. Some files are *valid UTF-8 containing double-encoded text* — content that
     was UTF-8, got misread as cp1252, and was re-saved as UTF-8, so "£" is
     literally stored as "Â£". Sniffing correctly reports UTF-8 and the mojibake
     survives into the text. `ftfy` repairs this; without it the garbage flows
     into every downstream embedding and LLM call.
"""
from __future__ import annotations

import io
import re
import shutil
import subprocess
import tempfile
from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt", ".html", ".htm", ".svg", ".xlsx", ".xls"}


class UnsupportedFileType(ValueError):
    pass


class ParseFailed(RuntimeError):
    pass


def _clean(text: str) -> str:
    """Repair mojibake, then normalize whitespace (collapse blank-line runs,
    strip trailing spaces). Applied to every parser's output, since double-encoded
    text can arrive via any format, not just .txt."""
    from ftfy import fix_text

    text = fix_text(text)
    lines = [ln.rstrip() for ln in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    out: list[str] = []
    blank = 0
    for ln in lines:
        if ln.strip():
            blank = 0
            out.append(ln)
        else:
            blank += 1
            if blank == 1:
                out.append("")
    return "\n".join(out).strip()


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------
def parse_txt(data: bytes) -> str:
    from charset_normalizer import from_bytes

    best = from_bytes(data).best()
    if best is None:
        # last resort: utf-8 with replacement rather than raising
        return _clean(data.decode("utf-8", errors="replace"))
    return _clean(str(best))


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------
def parse_html(data: bytes) -> str:
    return clean_html_text(data.decode("utf-8", errors="replace"))


# Cheap check before paying for a BeautifulSoup parse. Job-board exports commonly
# store descriptions as HTML fragments inside a spreadsheet cell, but plenty of
# sheets hold plain prose that merely mentions a "<" — require a closing tag too.
_HTML_MARKUP = re.compile(r"<(p|div|br|ul|ol|li|strong|em|b|i|span|h[1-6]|table)\b[^>]*>", re.I)


def looks_like_html(text: str) -> bool:
    return bool(_HTML_MARKUP.search(text))


def clean_html_text(text: str) -> str:
    """Strip markup from an HTML fragment, keeping block structure as newlines.

    Shared with spreadsheet ingestion: a job-board CSV routinely stores each
    description as ~5KB of HTML in one cell. Passing that to the LLM verbatim
    spends most of the token budget on tags and gives the model markup to reason
    about instead of prose.
    """
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(text, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    # get_text with a newline separator keeps block structure readable
    return _clean(soup.get_text(separator="\n"))


# ---------------------------------------------------------------------------
# PDF — pdfplumber primary (better layout handling), pypdf fallback
# ---------------------------------------------------------------------------
def parse_pdf(data: bytes) -> str:
    try:
        import pdfplumber

        pages: list[str] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        text = "\n\n".join(pages)
        if text.strip():
            return _clean(text)
    except Exception as e:  # noqa: BLE001 — fall through to pypdf
        print(f"  [parsers] pdfplumber failed ({e}); trying pypdf")

    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        text = "\n\n".join((p.extract_text() or "") for p in reader.pages)
        if not text.strip():
            raise ParseFailed("PDF produced no extractable text (likely a scanned image)")
        return _clean(text)
    except ParseFailed:
        raise
    except Exception as e:  # noqa: BLE001
        raise ParseFailed(f"could not extract text from PDF: {e}") from e


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def parse_docx(data: bytes) -> str:
    from docx import Document

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise ParseFailed(f"could not open DOCX: {e}") from e

    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return _clean("\n".join(parts))


# ---------------------------------------------------------------------------
# DOC (legacy binary) — no reliable pure-Python parser exists.
# Convert to .docx via headless LibreOffice, then parse as DOCX. This is the
# only approach that works on both Windows dev and a Linux container without
# depending on Word COM automation (Windows-only, and would break the
# deploy-to-container path).
# ---------------------------------------------------------------------------
def _find_soffice() -> str | None:
    for name in ("soffice", "soffice.exe", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found
    # common Windows install locations, which aren't usually on PATH
    for candidate in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ):
        if Path(candidate).exists():
            return candidate
    return None


def parse_doc(data: bytes) -> str:
    soffice = _find_soffice()
    if soffice is None:
        raise ParseFailed(
            "legacy .doc files need LibreOffice to convert to .docx first, and no "
            "soffice binary was found. Install LibreOffice, or ask the user to "
            "re-save the file as .docx / .pdf."
        )

    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        src = tmpdir / "input.doc"
        src.write_bytes(data)
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "docx", "--outdir", str(tmpdir), str(src)],
                check=True,
                capture_output=True,
                timeout=120,
            )
        except subprocess.CalledProcessError as e:
            raise ParseFailed(f"LibreOffice conversion failed: {e.stderr.decode(errors='replace')[:300]}") from e
        except subprocess.TimeoutExpired as e:
            raise ParseFailed("LibreOffice conversion timed out") from e

        converted = tmpdir / "input.docx"
        if not converted.exists():
            raise ParseFailed("LibreOffice reported success but produced no .docx")
        return parse_docx(converted.read_bytes())


# ---------------------------------------------------------------------------
# Process diagrams and spreadsheets — added for Workforce Studio step 2
# ---------------------------------------------------------------------------
def parse_svg(data: bytes) -> str:
    """Labels out of an SVG diagram, in document order.

    A process map drawn as SVG carries its step names as `<text>` nodes, so the labels
    are recoverable. The *arrows* are not: connector geometry would have to be matched
    back to boxes by coordinate, which is a different and much less reliable job. So
    this yields the labels and whatever ordering the document happens to have, and the
    UI says exactly that at upload rather than implying a reconstructed flowchart.

    `<tspan>` runs are joined without a separator because a single label is routinely
    split across them for line wrapping — "Approve" / "invoice" is one step, not two.
    """
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(data, "xml")
    if soup.find("svg") is None:
        # Fall back to the HTML parser: an SVG saved by a drawing tool is sometimes
        # wrapped in enough HTML that the XML parser finds no root svg.
        soup = BeautifulSoup(data, "lxml")

    lines: list[str] = []
    for node in soup.find_all(["text", "title", "desc"]):
        if node.name == "text":
            spans = node.find_all("tspan")
            text = "".join(s.get_text() for s in spans) if spans else node.get_text()
        else:
            text = node.get_text()
        text = re.sub(r"\s+", " ", text).strip()
        if text and text not in lines[-1:]:
            lines.append(text)
    return _clean("\n".join(lines))


def parse_xlsx(data: bytes) -> str:
    """Every sheet as tab-separated rows, sheet names as headings.

    Deliberately not a table-shape guess. A process or catalogue spreadsheet may put
    its header on row 4, use merged cells, or spread across sheets; inferring a schema
    here would be a second column-mapping problem. The model reads the text and works
    it out, which is what it is good at.
    """
    import pandas as pd

    try:
        sheets = pd.read_excel(io.BytesIO(data), sheet_name=None, header=None, dtype=str)
    except Exception as e:  # noqa: BLE001 — pandas raises a wide variety here
        raise ParseFailed(f"could not read the spreadsheet: {type(e).__name__}: {e}") from e

    parts: list[str] = []
    for name, df in sheets.items():
        rows = [
            "\t".join(str(v) for v in row if str(v) not in ("nan", "None", ""))
            for row in df.fillna("").itertuples(index=False)
        ]
        rows = [r for r in rows if r.strip()]
        if rows:
            parts.append(f"## {name}\n" + "\n".join(rows))
    return _clean("\n\n".join(parts))


# ---------------------------------------------------------------------------
# Dispatch
# ---------------------------------------------------------------------------
_PARSERS = {
    ".txt": parse_txt,
    ".html": parse_html,
    ".htm": parse_html,
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".doc": parse_doc,
    ".svg": parse_svg,
    ".xlsx": parse_xlsx,
    ".xls": parse_xlsx,
}


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded job-description file, by extension."""
    ext = Path(filename).suffix.lower()
    parser = _PARSERS.get(ext)
    if parser is None:
        raise UnsupportedFileType(
            f"unsupported file type '{ext}' — supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    text = parser(data)
    if not text.strip():
        raise ParseFailed(f"no text could be extracted from '{filename}'")
    return text
