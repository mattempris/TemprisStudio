"""Job profile export to PDF and DocX.

instructions.txt: "beautiful job profile documents (html native, PDF/DocX export)".

Two independent renderers off the one canonical structured content model, NOT one
converter serving both:

  PDF  — WeasyPrint over the same HTML/CSS the app displays, so the PDF is a
         faithful reproduction of the on-screen document.
  DocX — python-docx building a genuinely native Word document (real heading
         styles, real bullet lists, a real table) directly from the structured
         content. This is a hard requirement from the user: HR/comp teams must be
         able to open it in Word and edit it with native styles. HTML-to-DocX
         converters produce a mangled approximation with the card CSS baked into
         broken tables, which is not editable in any meaningful sense.
"""
from __future__ import annotations

import io


class PdfExportUnavailable(RuntimeError):
    """WeasyPrint's native deps (Pango/Cairo, via GTK3 on Windows) are missing."""


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def render_pdf(html: str, *, base_url: str | None = None) -> bytes:
    """Render the profile HTML to PDF.

    On Linux (including the app's container) WeasyPrint's deps come from
    `libpango`/`libcairo`, installed in the Dockerfile. On native Windows they
    come from the separately-installed GTK3 runtime, which is often absent — so
    the import failure is translated into an actionable error rather than a
    stack trace about a missing DLL.
    """
    try:
        from weasyprint import HTML
    except (ImportError, OSError) as e:
        raise PdfExportUnavailable(
            "PDF export needs WeasyPrint's native libraries (Pango/Cairo). On "
            "Windows these come from the GTK3 runtime; install it, or run PDF "
            f"export via the Linux backend container. Underlying error: {e}"
        ) from e

    return HTML(string=html, base_url=base_url).write_pdf()


def pdf_available() -> bool:
    try:
        import weasyprint  # noqa: F401

        return True
    except (ImportError, OSError):
        return False


# ---------------------------------------------------------------------------
# DocX — native Word document, built from structured content
# ---------------------------------------------------------------------------
def render_docx(
    content: dict,
    *,
    company_name: str | None = None,
    job_level: str | None = None,
    about_company: str | None = None,
    diversity_statement: str | None = None,
    accent_color: str = "#1d4ed8",
) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Restyle the built-in styles rather than applying direct formatting, so the
    # document stays properly editable: changing "Heading 1" in Word updates
    # every heading, as a Word user expects.
    accent_rgb = _hex_to_rgbcolor(accent_color, RGBColor)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for style_name, size, bold in (("Heading 1", 20, True), ("Heading 2", 13, True), ("Heading 3", 11, True)):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = accent_rgb

    # ---- title block ----
    doc.add_heading(content.get("title", "Job Profile"), level=1)

    subtitle_bits = [
        b
        for b in (company_name, content.get("family"), content.get("category"), content.get("level_context"))
        if b
    ]
    if subtitle_bits:
        p = doc.add_paragraph(" · ".join(subtitle_bits))
        p.runs[0].font.size = Pt(9.5)
        p.runs[0].font.color.rgb = RGBColor(0x4A, 0x4F, 0x63)

    if content.get("badges"):
        p = doc.add_paragraph(" | ".join(content["badges"]))
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(0x7A, 0x80, 0x95)

    if job_level:
        p = doc.add_paragraph()
        run = p.add_run(f"Evaluated job level: {job_level}")
        run.bold = True
        run.font.size = Pt(10)

    # ---- body sections ----
    if about_company:
        doc.add_heading(f"About {company_name or 'the Organisation'}", level=2)
        doc.add_paragraph(about_company)

    if content.get("about_role"):
        doc.add_heading("About the Role", level=2)
        for para in content["about_role"]:
            doc.add_paragraph(para)

    # "About You" uses a real 2-column Word table, mirroring the HTML's two-col
    # grid while remaining a native, editable Word table.
    has_left = bool(content.get("requirements") or content.get("essential_skills"))
    has_right = bool(content.get("desirable_skills"))
    if has_left or has_right:
        doc.add_heading("About You", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        left, right = table.rows[0].cells

        _fill_cell(left, [("Minimum Requirements", content.get("requirements")),
                          ("Essential Skills", content.get("essential_skills"))], Pt)
        _fill_cell(right, [("Desirable Skills", content.get("desirable_skills"))], Pt)

        if content.get("tags"):
            p = right.add_paragraph(", ".join(content["tags"]))
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.italic = True

    if content.get("responsibilities"):
        doc.add_heading("Key Responsibilities", level=2)
        for item in content["responsibilities"]:
            doc.add_paragraph(item, style="List Bullet")

    if content.get("contribution"):
        doc.add_heading("Your Contribution", level=2)
        doc.add_paragraph("You can expect to engage in:")
        for item in content["contribution"]:
            doc.add_paragraph(item, style="List Bullet")

    if content.get("required_of_you"):
        doc.add_heading("Required of You", level=2)
        for item in content["required_of_you"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{item.get('label', '')}: ").bold = True
            p.add_run(str(item.get("value", "")))

    if content.get("reporting_line") or content.get("budget_responsibility"):
        doc.add_heading("Scope & Accountability", level=2)
        if content.get("reporting_line"):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("Reporting line: ").bold = True
            p.add_run(content["reporting_line"])
        if content.get("budget_responsibility"):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("Budget responsibility: ").bold = True
            p.add_run(content["budget_responsibility"])

    if diversity_statement:
        doc.add_heading("Diversity Statement", level=2)
        doc.add_paragraph(diversity_statement)

    footer = doc.add_paragraph(
        "This job profile describes the general scope of the role and is not an "
        "exhaustive list of duties."
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer.runs[0].font.size = Pt(8.5)
    footer.runs[0].font.color.rgb = RGBColor(0x7A, 0x80, 0x95)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _fill_cell(cell, sections: list[tuple[str, list | None]], Pt) -> None:
    """Write heading+bullets pairs into a table cell, dropping the empty
    paragraph python-docx creates with every new cell."""
    first = True
    for heading, items in sections:
        if not items:
            continue
        if first and cell.paragraphs and not cell.paragraphs[0].text:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        run = p.add_run(heading)
        run.bold = True
        run.font.size = Pt(10)
        first = False
        for item in items:
            bullet = cell.add_paragraph(style="List Bullet")
            bullet.add_run(item).font.size = Pt(9.5)


def _hex_to_rgbcolor(hex_color: str, RGBColor):
    h = hex_color.lstrip("#")
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    try:
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    except (ValueError, IndexError):
        return RGBColor(0x1D, 0x4E, 0xD8)
